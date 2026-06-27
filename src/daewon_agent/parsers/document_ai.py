"""Step 1: Google Cloud Document AI Layout Parser 로 비정형 문서 파싱.

물성 데이터(.docx) 및 DMF(.pdf) 의 구조를 인식해 평문 텍스트로 변환한다.
.docx 는 Document AI 가 직접 지원하지 않으므로 python-docx 로 폴백한다.
"""
from __future__ import annotations

import logging
from pathlib import Path

from google.api_core.client_options import ClientOptions
from google.cloud import documentai

from ..config import Settings

logger = logging.getLogger(__name__)

# Document AI 가 직접 처리 가능한 MIME 타입
_MIME = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".tiff": "image/tiff",
    ".png": "image/png",
    ".jpg": "image/jpeg",
}


class DocumentParser:
    def __init__(self, settings: Settings):
        self.settings = settings
        loc = settings.document_ai.processor_location
        self._client = documentai.DocumentProcessorServiceClient(
            client_options=ClientOptions(
                api_endpoint=f"{loc}-documentai.googleapis.com"
            )
        )
        self._processor_name = self._client.processor_path(
            settings.gcp.project_id,
            loc,
            settings.document_ai.processor_id,
        )

    def parse(self, file_path: str | Path) -> str:
        """파일을 평문 텍스트로 파싱."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix not in _MIME:
            raise ValueError(f"지원하지 않는 형식입니다: {suffix}")

        return self._parse_with_document_ai(path, _MIME[suffix])

    def _parse_with_document_ai(self, path: Path, mime_type: str) -> str:
        content = path.read_bytes()
        raw_document = documentai.RawDocument(content=content, mime_type=mime_type)
        request = documentai.ProcessRequest(
            name=self._processor_name, raw_document=raw_document
        )
        logger.info("Document AI 파싱: %s", path.name)
        result = self._client.process_document(request=request)
        return result.document.text

    @staticmethod
    def _parse_docx(path: Path) -> str:
        """.docx 폴백 파서 (표 포함)."""
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n".join(parts)
