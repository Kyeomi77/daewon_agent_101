"""Step 7: 최적 처방 가이드 .docx 자동 생성.

확정된 처방과 편차 분석 결과를 종합해 최종 보고서를 생성한다.
사내 문서 관리 시스템(클라우독/ELN) 연동은 eln 모듈에서 처리한다.
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..schemas import FormulationProposal, DeviationAnalysis

logger = logging.getLogger(__name__)


def write_guide_docx(
    proposal: FormulationProposal,
    analyses: list[DeviationAnalysis] | None,
    out_path: str | Path,
) -> Path:
    out_path = Path(out_path)
    doc = Document()

    title = doc.add_heading("최적 처방 가이드", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"API: {proposal.api_name}    "
        f"BCS: {proposal.bcs_class.value if proposal.bcs_class else '미상'}    "
        f"생성일: {datetime.now():%Y-%m-%d %H:%M}"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    analysis_map = {a.formulation_id: a for a in (analyses or [])}

    doc.add_heading("1. 처방 요약", level=1)
    for f in proposal.formulations:
        a = analysis_map.get(f.formulation_id)
        status = ""
        if a is not None:
            status = "  [확정 ✔]" if a.passed else "  [개선 필요]"
        doc.add_heading(f"처방 {f.formulation_id} (리스크: {f.risk_level.value}){status}",
                        level=2)

        doc.add_paragraph(f.rationale)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        for i, h in enumerate(["부형제", "역할", "함량(mg)", "비율(%)"]):
            hdr[i].paragraphs[0].add_run(h).bold = True
        for e in f.excipients:
            cells = table.add_row().cells
            cells[0].text = e.name
            cells[1].text = e.function
            cells[2].text = str(e.amount_mg)
            cells[3].text = str(e.percent) if e.percent is not None else "-"

        p = doc.add_paragraph()
        p.add_run(
            f"설계 기대 — 용출률 {f.expected_dissolution}%, 함량 {f.expected_assay}%"
        ).bold = True

        if a is not None:
            dev = doc.add_paragraph()
            dev.add_run(
                f"편차 분석 — Δ용출 {a.delta_dissolution}%, Δ함량 {a.delta_assay}% "
                f"→ {'합격' if a.passed else '불합격'}"
            )
            if a.improvement_plan:
                doc.add_heading("개선안", level=3)
                doc.add_paragraph(a.improvement_plan)

    if proposal.cited_cases:
        doc.add_heading("2. 참조 유사 사례", level=1)
        doc.add_paragraph(", ".join(proposal.cited_cases))

    doc.add_heading("3. 감사 추적 (Audit Trail)", level=1)
    doc.add_paragraph(
        "본 가이드는 Vertex AI 기반 제제연구 에이전트가 자동 생성하였으며, "
        "GxP 규제 준수를 위해 입력 파라미터·추론 모델·합격 기준이 로그에 기록됩니다."
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    logger.info("최종 처방 가이드 저장: %s", out_path)
    return out_path
